"""
Turn an uploaded enterprise document into (heading, body) sections.

This is intentionally heuristic rather than a full document-structure parser:
- DOCX: use paragraph "Heading N" styles, which most enterprise docs already use.
- PDF: use font size deltas (pdfplumber) as a heading signal, falling back to
  a simple short-line heuristic if font info isn't usable.
- TXT/MD: use '#'-style markdown headings, falling back to blank-line paragraphs.

Good enough to preserve real document hierarchy without building a full
layout-analysis engine. YAGNI.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass
class Section:
    heading: str
    level: int
    body: str


def parse_document(path: Path) -> list[Section]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    return _parse_text(path)


# ---------- DOCX ----------

def _parse_docx(path: Path) -> list[Section]:
    import docx  # python-docx

    doc = docx.Document(str(path))
    sections: list[Section] = []
    current_heading = "Introduction"
    current_level = 1
    current_body: list[str] = []

    def flush():
        text = "\n".join(current_body).strip()
        if text:
            sections.append(Section(current_heading, current_level, text))

    for para in doc.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            flush()
            current_body = []
            current_heading = text
            try:
                current_level = int(style.replace("heading", "").strip() or 1)
            except ValueError:
                current_level = 1
        else:
            current_body.append(text)
    flush()

    # also pull table content in as plain text appended to the last section
    if doc.tables:
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
        if table_lines:
            sections.append(Section("Tables", 1, "\n".join(table_lines)))

    return [s for s in sections if s.body or s.heading]


# ---------- PDF ----------

def _parse_pdf(path: Path) -> list[Section]:
    import pdfplumber

    lines_with_size: list[tuple[str, float]] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            chars = page.chars
            if not chars:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    if line.strip():
                        lines_with_size.append((line.strip(), 0.0))
                continue
            # group chars into lines by 'top' position
            lines: dict[float, list] = {}
            for ch in chars:
                key = round(ch["top"], 1)
                lines.setdefault(key, []).append(ch)
            for top in sorted(lines.keys()):
                chs = sorted(lines[top], key=lambda c: c.get("x0", 0.0))
                line_parts = []
                for idx, c in enumerate(chs):
                    line_parts.append(c["text"])
                    if idx < len(chs) - 1:
                        next_c = chs[idx + 1]
                        gap = next_c.get("x0", 0.0) - c.get("x1", 0.0)
                        if gap > 2.0 and c["text"] != " " and next_c["text"] != " ":
                            line_parts.append(" ")
                text = "".join(line_parts).strip()
                if text:
                    avg_size = sum(c.get("size", 10.0) for c in chs) / len(chs)
                    lines_with_size.append((text, avg_size))

    if not lines_with_size:
        return [Section("Document", 1, "")]

    sizes = [s for _, s in lines_with_size if s > 0]
    body_size = _mode(sizes) if sizes else 0

    sections: list[Section] = []
    current_heading = "Introduction"
    current_body: list[str] = []

    def flush():
        text = "\n".join(current_body).strip()
        if text or not sections:
            sections.append(Section(current_heading, 1, text))

    for text, size in lines_with_size:
        is_heading = (
            size > 0
            and body_size > 0
            and size >= body_size * 1.15
            and len(text) < 120
        )
        if is_heading:
            flush()
            current_body = []
            current_heading = text
        else:
            current_body.append(text)
    flush()

    return [s for s in sections if s.body or s.heading]


def _mode(values: list[float]) -> float:
    from collections import Counter

    rounded = [round(v, 1) for v in values]
    return Counter(rounded).most_common(1)[0][0]


# ---------- TXT / MD ----------

def _parse_text(path: Path) -> list[Section]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    if any(line.strip().startswith("#") for line in lines):
        sections: list[Section] = []
        current_heading = "Introduction"
        current_level = 1
        current_body: list[str] = []

        def flush():
            body = "\n".join(current_body).strip()
            if body or not sections:
                sections.append(Section(current_heading, current_level, body))

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                flush()
                current_body = []
                level = len(stripped) - len(stripped.lstrip("#"))
                current_heading = stripped.lstrip("#").strip()
                current_level = level
            else:
                current_body.append(line)
        flush()
        return [s for s in sections if s.body or s.heading]

    # fallback: split on blank lines, first line of each paragraph is the heading
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    sections = []
    for para in paragraphs:
        plines = para.splitlines()
        heading = plines[0][:80]
        body = "\n".join(plines[1:]).strip() or plines[0]
        sections.append(Section(heading, 1, body))
    return sections or [Section("Document", 1, text)]
